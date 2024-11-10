import os
import datetime
from docx import Document
import fitz  # PyMuPDF
import pytesseract
from PIL import Image
import io

# Paths
pdf_folder = "pdf"
word_folder = "word"
extracted_log = "extracted.txt"
today = datetime.datetime.now().strftime("%Y-%m-%d")
output_word_file = os.path.join(word_folder, f"extracted_{today}.docx")

# Read previously extracted PDF names from the log file
if os.path.exists(extracted_log):
    with open(extracted_log, "r") as file:
        extracted_files = set(file.read().splitlines())
else:
    extracted_files = set()

# Initialize Word document
doc = Document()
doc.add_heading(f"Extracted PDF Content - {today}", level=1)

# Function to extract text from regular PDFs
def extract_text_from_pdf(pdf_path):
    doc_text = ""
    pdf_doc = fitz.open(pdf_path)
    for page in pdf_doc:
        doc_text += page.get_text()
    pdf_doc.close()
    return doc_text

# Function to extract text from image-based PDFs using OCR
def extract_text_with_ocr(pdf_path):
    pdf_doc = fitz.open(pdf_path)
    ocr_text = ""
    for page in pdf_doc:
        # Extract image from each page
        image_list = page.get_images(full=True)
        for img_index, img in enumerate(image_list):
            xref = img[0]
            base_image = pdf_doc.extract_image(xref)
            image_bytes = base_image["image"]
            
            # Convert image bytes to an Image object and apply OCR
            image = Image.open(io.BytesIO(image_bytes))
            ocr_text += pytesseract.image_to_string(image) + "\n"
    pdf_doc.close()
    return ocr_text

# Process each PDF in the pdf folder
new_files = []
for pdf_file in os.listdir(pdf_folder):
    pdf_path = os.path.join(pdf_folder, pdf_file)
    if pdf_file.endswith(".pdf") and pdf_file not in extracted_files:
        # Extract text, try regular extraction first, then OCR if empty
        text = extract_text_from_pdf(pdf_path)
        if not text.strip():  # If no text was found, apply OCR
            text = extract_text_with_ocr(pdf_path)

        # Add extracted content to the Word document
        doc.add_heading(pdf_file, level=2)
        doc.add_paragraph(text)
        
        # Log this PDF as processed
        new_files.append(pdf_file)
        extracted_files.add(pdf_file)

# Save the Word document if there were new files
if new_files:
    doc.save(output_word_file)
    with open(extracted_log, "a") as file:
        for pdf_file in new_files:
            file.write(pdf_file + "\n")

    print(f"Processed and saved new PDFs into: {output_word_file}")
    print(f"Updated log file: {extracted_log}")
else:
    print("No new PDFs found to process.")
